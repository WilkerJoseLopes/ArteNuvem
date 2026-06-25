import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Define a cor de fundo de uma célula usando XML."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Define as margens internas (padding) de uma célula em dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def main():
    doc = Document()

    # Configuração de Estilo Global (Normal)
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Título do Documento
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run_title = title.add_run("Dicionário de Dados — Tabela 'notification'")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Azul Marinho Escuro

    # Subtítulo/Descrição Breve
    desc = doc.add_paragraph()
    desc.paragraph_format.space_after = Pt(18)
    run_desc = desc.add_run(
        "Este documento descreve a estrutura física da tabela 'notification' (Notificações) "
        "conforme definida na base de dados do ArteNuvem (s.sql)."
    )
    run_desc.font.size = Pt(10)
    run_desc.font.italic = True
    run_desc.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Dados da tabela
    headers = ["Campo", "Tipo", "Descrição", "Restrições"]
    rows_data = [
        [
            "ID_Notification",
            "INTEGER",
            "Identificador único e sequencial de cada notificação gerada.",
            "PRIMARY KEY (PK), NOT NULL, AUTO-INCREMENT (seq)"
        ],
        [
            "UserID",
            "INTEGER",
            "Chave estrangeira que associa a notificação ao utilizador de destino.",
            "FOREIGN KEY (FK) -> public.utilizador(ID_Utilizador), NOT NULL"
        ],
        [
            "Type",
            "VARCHAR",
            "Tipo de interação que motivou a notificação (ex: 'like' ou 'comentario').",
            "NOT NULL"
        ],
        [
            "ID_Imagem",
            "INTEGER",
            "Referência à obra/imagem associada à notificação (opcional).",
            "FOREIGN KEY (FK) -> public.imagem(ID_Imagem), NULL"
        ],
        [
            "Message",
            "VARCHAR",
            "Texto descritivo da notificação a ser exibido na interface do utilizador.",
            "NOT NULL"
        ],
        [
            "Count",
            "INTEGER",
            "Contador para agrupar notificações semelhantes de forma consolidada.",
            "NOT NULL, DEFAULT 1"
        ],
        [
            "IsRead",
            "BOOLEAN",
            "Flag indicando se a notificação já foi lida (TRUE) ou permanece não lida (FALSE).",
            "NOT NULL, DEFAULT FALSE"
        ],
        [
            "CreatedAt",
            "TIMESTAMP",
            "Data e hora em que a notificação foi inserida no sistema.",
            "NOT NULL, DEFAULT NOW()"
        ]
    ]

    # Criar a Tabela
    # Rows = data + 1 (header)
    table = doc.add_table(rows=len(rows_data) + 1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Definir larguras de colunas
    col_widths = [Inches(1.4), Inches(1.1), Inches(3.2), Inches(1.8)]

    # Cabeçalho da Tabela
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = col_widths[i]
        set_cell_background(cell, "1E3A8A") # Azul Marinho Escuro
        set_cell_margins(cell, top=160, bottom=160, left=180, right=180) # Maior padding no header
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title_text)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # Texto Branco

    # Preencher Dados
    for r_idx, row_content in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        # Alternância de cores (zebra striping)
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, cell_text in enumerate(row_content):
            cell = row_cells[c_idx]
            cell.width = col_widths[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            
            # Destacar nomes dos campos em negrito
            if c_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            # Destacar restrições
            elif c_idx == 3:
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    output_path = r"c:\Users\lopes\Music\SASASA\ArteNuvem-main (2) (1)\ArteNuvem-main\Dicionario_Dados_Notificacoes.docx"
    doc.save(output_path)
    print(f"Documento Word criado com sucesso em: {output_path}")

if __name__ == "__main__":
    main()
